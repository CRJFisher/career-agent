"""Document parsing utilities for extracting text from various formats.

This module provides parsers for PDF, Word, Markdown, and Google Docs,
returning structured content with metadata.
"""

from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
import re
import logging
import PyPDF2
from docx import Document
import markdown
import frontmatter
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
import tempfile
import os


logger = logging.getLogger(__name__)


@dataclass
class DocumentSection:
    """Represents a section within a document."""
    
    title: str
    content: str
    level: int  # Heading level (1-6)
    start_page: Optional[int] = None
    end_page: Optional[int] = None
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            'title': self.title,
            'content': self.content,
            'level': self.level,
            'start_page': self.start_page,
            'end_page': self.end_page
        }


@dataclass
class ParsedDocument:
    """Standardized output from document parsers."""
    
    content: str  # Full text content
    metadata: Dict[str, Any] = field(default_factory=dict)
    sections: List[DocumentSection] = field(default_factory=list)
    source_type: str = ""  # pdf, markdown, docx, gdoc
    source_path: str = ""
    raw_data: Optional[Any] = None  # Original parsed data
    error: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for serialization."""
        return {
            'content': self.content,
            'metadata': self.metadata,
            'sections': [s.to_dict() for s in self.sections],
            'source_type': self.source_type,
            'source_path': self.source_path,
            'error': self.error
        }
    
    def get_text(self) -> str:
        """Get all text content."""
        return self.content
    
    def get_sections_by_level(self, level: int) -> List[DocumentSection]:
        """Get all sections at a specific heading level."""
        return [s for s in self.sections if s.level == level]
    
    def search_content(self, pattern: str, case_sensitive: bool = False) -> List[Dict[str, Any]]:
        """Search for pattern in content."""
        flags = 0 if case_sensitive else re.IGNORECASE
        matches = []
        
        for match in re.finditer(pattern, self.content, flags):
            # Find which section contains this match
            section = None
            for s in self.sections:
                if match.start() >= self.content.find(s.content):
                    section = s
                else:
                    break
            
            matches.append({
                'text': match.group(),
                'start': match.start(),
                'end': match.end(),
                'section': section.title if section else None
            })
        
        return matches


class DocumentParser(ABC):
    """Abstract base class for document parsers."""
    
    def __init__(self):
        self.extract_metadata = True
        self.extract_sections = True
    
    @abstractmethod
    def can_parse(self, file_path: Union[str, Path]) -> bool:
        """Check if this parser can handle the file."""
        pass
    
    @abstractmethod
    def parse(self, file_path: Union[str, Path]) -> ParsedDocument:
        """Parse document and return structured content."""
        pass
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove multiple spaces
        text = re.sub(r'\s+', ' ', text)
        # Remove zero-width characters
        text = re.sub(r'[\u200b\u200c\u200d\ufeff]', '', text)
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        # Remove excessive line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()


class PDFParser(DocumentParser):
    """Parser for PDF documents."""
    
    def can_parse(self, file_path: Union[str, Path]) -> bool:
        """Check if file is a PDF."""
        path = Path(file_path)
        return path.suffix.lower() == '.pdf' and path.exists()
    
    def parse(self, file_path: Union[str, Path]) -> ParsedDocument:
        """Extract text and structure from PDF."""
        path = Path(file_path)
        doc = ParsedDocument(
            content='',  # Initialize with empty content
            source_type='pdf',
            source_path=str(path.absolute())
        )
        
        try:
            with open(path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                if self.extract_metadata:
                    doc.metadata = self._extract_pdf_metadata(pdf_reader)
                
                # Extract text content
                all_text = []
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text:
                        all_text.append(text)
                    
                    # Try to detect sections (basic heuristic)
                    if self.extract_sections:
                        sections = self._extract_pdf_sections(text, page_num + 1)
                        doc.sections.extend(sections)
                
                doc.content = self._clean_text('\n'.join(all_text))
                
        except Exception as e:
            logger.error(f"Error parsing PDF {path}: {e}")
            doc.error = str(e)
        
        return doc
    
    def _extract_pdf_metadata(self, pdf_reader: PyPDF2.PdfReader) -> Dict[str, Any]:
        """Extract metadata from PDF."""
        metadata = {}
        
        if pdf_reader.metadata:
            info = pdf_reader.metadata
            
            # Standard PDF metadata fields
            fields = ['/Title', '/Author', '/Subject', '/Creator', 
                     '/Producer', '/CreationDate', '/ModDate']
            
            for field in fields:
                if field in info:
                    key = field[1:].lower()  # Remove '/' and lowercase
                    value = info[field]
                    
                    # Convert dates
                    if 'date' in key.lower() and hasattr(value, 'year'):
                        metadata[key] = value.strftime('%Y-%m-%d %H:%M:%S')
                    else:
                        metadata[key] = str(value)
        
        metadata['page_count'] = len(pdf_reader.pages)
        
        return metadata
    
    def _extract_pdf_sections(self, text: str, page_num: int) -> List[DocumentSection]:
        """Extract sections from PDF text (basic heuristic)."""
        sections = []
        
        # Look for common heading patterns
        patterns = [
            (r'^([A-Z][A-Z\s]+)$', 1),  # ALL CAPS lines
            (r'^(\d+\.?\s+[A-Z][^.]+)$', 2),  # Numbered headings
            (r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*):?\s*$', 3),  # Title Case
        ]
        
        lines = text.split('\n')
        for i, line in enumerate(lines):
            line = line.strip()
            if not line or len(line) > 100:  # Skip empty or very long lines
                continue
            
            for pattern, level in patterns:
                if re.match(pattern, line):
                    # Get content until next heading
                    content_lines = []
                    for j in range(i + 1, len(lines)):
                        next_line = lines[j].strip()
                        if any(re.match(p[0], next_line) for p in patterns):
                            break
                        content_lines.append(lines[j])
                    
                    sections.append(DocumentSection(
                        title=line,
                        content='\n'.join(content_lines).strip(),
                        level=level,
                        start_page=page_num,
                        end_page=page_num
                    ))
                    break
        
        return sections


class MarkdownParser(DocumentParser):
    """Parser for Markdown documents."""
    
    def can_parse(self, file_path: Union[str, Path]) -> bool:
        """Check if file is Markdown."""
        path = Path(file_path)
        return path.suffix.lower() in ['.md', '.markdown'] and path.exists()
    
    def parse(self, file_path: Union[str, Path]) -> ParsedDocument:
        """Parse Markdown with frontmatter support."""
        path = Path(file_path)
        doc = ParsedDocument(
            content='',  # Initialize with empty content
            source_type='markdown',
            source_path=str(path.absolute())
        )
        
        try:
            # Parse with frontmatter
            with open(path, 'r', encoding='utf-8') as f:
                post = frontmatter.load(f)
            
            # Extract metadata from frontmatter
            if self.extract_metadata:
                doc.metadata = {}
                for key, value in post.metadata.items():
                    # Convert dates to strings
                    if hasattr(value, 'isoformat'):
                        doc.metadata[key] = value.isoformat()
                    elif hasattr(value, 'strftime'):
                        doc.metadata[key] = str(value)
                    else:
                        doc.metadata[key] = value
                
                # Add file metadata
                stat = path.stat()
                doc.metadata['file_modified'] = datetime.fromtimestamp(stat.st_mtime).isoformat()
                doc.metadata['file_size'] = stat.st_size
            
            # Parse content (preserve markdown structure)
            content = post.content
            doc.content = content  # Don't clean markdown - preserve formatting
            
            # Extract sections from headings
            if self.extract_sections:
                doc.sections = self._extract_markdown_sections(content)
            
            # Store both raw and HTML versions
            doc.raw_data = {
                'markdown': content,
                'html': markdown.markdown(content, extensions=['extra', 'toc'])
            }
            
        except Exception as e:
            logger.error(f"Error parsing Markdown {path}: {e}")
            doc.error = str(e)
        
        return doc
    
    def _extract_markdown_sections(self, content: str) -> List[DocumentSection]:
        """Extract sections from Markdown headings."""
        sections = []
        lines = content.split('\n')
        
        current_section = None
        current_content = []
        
        for line in lines:
            # Check for heading
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            
            if heading_match:
                # Save previous section
                if current_section:
                    current_section.content = self._clean_text('\n'.join(current_content))
                    sections.append(current_section)
                
                # Start new section
                level = len(heading_match.group(1))
                title = heading_match.group(2).strip()
                current_section = DocumentSection(
                    title=title,
                    content='',
                    level=level
                )
                current_content = []
            else:
                current_content.append(line)
        
        # Save last section
        if current_section:
            current_section.content = self._clean_text('\n'.join(current_content))
            sections.append(current_section)
        
        return sections


class DocxParser(DocumentParser):
    """Parser for Word documents."""
    
    def can_parse(self, file_path: Union[str, Path]) -> bool:
        """Check if file is a Word document."""
        path = Path(file_path)
        return path.suffix.lower() in ['.docx', '.doc'] and path.exists()
    
    def parse(self, file_path: Union[str, Path]) -> ParsedDocument:
        """Extract text and structure from Word document."""
        path = Path(file_path)
        doc = ParsedDocument(
            content='',  # Initialize with empty content
            source_type='docx',
            source_path=str(path.absolute())
        )
        
        try:
            # Note: python-docx only supports .docx, not .doc
            if path.suffix.lower() == '.doc':
                doc.error = "Legacy .doc format not supported. Please convert to .docx"
                return doc
            
            docx = Document(path)
            
            # Extract metadata
            if self.extract_metadata:
                doc.metadata = self._extract_docx_metadata(docx)
            
            # Extract content
            all_text = []
            current_section = None
            section_content = []
            
            for paragraph in docx.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                all_text.append(text)
                
                # Check if this is a heading
                if self.extract_sections and paragraph.style.name.startswith('Heading'):
                    # Save previous section
                    if current_section:
                        current_section.content = self._clean_text('\n'.join(section_content))
                        doc.sections.append(current_section)
                    
                    # Extract heading level
                    level_match = re.search(r'Heading (\d)', paragraph.style.name)
                    level = int(level_match.group(1)) if level_match else 1
                    
                    # Start new section
                    current_section = DocumentSection(
                        title=text,
                        content='',
                        level=level
                    )
                    section_content = []
                else:
                    section_content.append(text)
            
            # Save last section
            if current_section:
                current_section.content = self._clean_text('\n'.join(section_content))
                doc.sections.append(current_section)
            
            doc.content = self._clean_text('\n'.join(all_text))
            
            # Extract tables
            doc.raw_data = {
                'tables': self._extract_docx_tables(docx)
            }
            
        except Exception as e:
            logger.error(f"Error parsing Word document {path}: {e}")
            doc.error = str(e)
        
        return doc
    
    def _extract_docx_metadata(self, docx: Document) -> Dict[str, Any]:
        """Extract metadata from Word document."""
        metadata = {}
        
        # Core properties
        props = docx.core_properties
        if props.author:
            metadata['author'] = props.author
        if props.title:
            metadata['title'] = props.title
        if props.subject:
            metadata['subject'] = props.subject
        if props.created:
            metadata['created'] = props.created.isoformat()
        if props.modified:
            metadata['modified'] = props.modified.isoformat()
        
        # Document statistics
        metadata['paragraph_count'] = len(docx.paragraphs)
        metadata['table_count'] = len(docx.tables)
        
        return metadata
    
    def _extract_docx_tables(self, docx: Document) -> List[List[List[str]]]:
        """Extract tables from Word document."""
        tables = []
        
        for table in docx.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables.append(table_data)
        
        return tables


class GoogleDocsParser(DocumentParser):
    """Parser for Google Docs (requires authentication)."""
    
    def __init__(self, service=None):
        super().__init__()
        self.service = service
    
    def can_parse(self, file_path: Union[str, Path]) -> bool:
        """Check if this is a Google Docs ID."""
        # Simple check - Google Doc IDs are typically alphanumeric strings
        doc_id = str(file_path)
        return bool(re.match(r'^[a-zA-Z0-9_-]+$', doc_id)) and len(doc_id) > 20
    
    def parse(self, file_path: Union[str, Path]) -> ParsedDocument:
        """Extract content from Google Doc."""
        doc_id = str(file_path)
        doc = ParsedDocument(
            content='',  # Initialize with empty content
            source_type='gdoc',
            source_path=doc_id
        )
        
        if not self.service:
            doc.error = "Google Docs service not initialized"
            return doc
        
        try:
            # Get document
            gdoc = self.service.documents().get(documentId=doc_id).execute()
            
            # Extract metadata
            if self.extract_metadata:
                doc.metadata = {
                    'title': gdoc.get('title', ''),
                    'documentId': doc_id
                }
            
            # Extract content
            content = self._extract_gdoc_content(gdoc)
            doc.content = self._clean_text(content)
            
            # Extract sections
            if self.extract_sections:
                doc.sections = self._extract_gdoc_sections(gdoc)
            
            doc.raw_data = gdoc
            
        except HttpError as e:
            logger.error(f"Error accessing Google Doc {doc_id}: {e}")
            doc.error = str(e)
        except Exception as e:
            logger.error(f"Error parsing Google Doc {doc_id}: {e}")
            doc.error = str(e)
        
        return doc
    
    def _extract_gdoc_content(self, gdoc: dict) -> str:
        """Extract text content from Google Doc structure."""
        content = []
        
        for element in gdoc.get('body', {}).get('content', []):
            if 'paragraph' in element:
                paragraph = element['paragraph']
                for elem in paragraph.get('elements', []):
                    if 'textRun' in elem:
                        content.append(elem['textRun'].get('content', ''))
        
        return ''.join(content)
    
    def _extract_gdoc_sections(self, gdoc: dict) -> List[DocumentSection]:
        """Extract sections from Google Doc headings."""
        sections = []
        current_section = None
        section_content = []
        
        for element in gdoc.get('body', {}).get('content', []):
            if 'paragraph' in element:
                paragraph = element['paragraph']
                style = paragraph.get('paragraphStyle', {}).get('namedStyleType', '')
                
                # Get text
                text_elements = []
                for elem in paragraph.get('elements', []):
                    if 'textRun' in elem:
                        text_elements.append(elem['textRun'].get('content', ''))
                text = ''.join(text_elements).strip()
                
                if style.startswith('HEADING_'):
                    # Save previous section
                    if current_section:
                        current_section.content = self._clean_text(''.join(section_content))
                        sections.append(current_section)
                    
                    # Extract level
                    level = int(style.replace('HEADING_', ''))
                    
                    # Start new section
                    current_section = DocumentSection(
                        title=text,
                        content='',
                        level=level
                    )
                    section_content = []
                else:
                    section_content.append(text)
        
        # Save last section
        if current_section:
            current_section.content = self._clean_text(''.join(section_content))
            sections.append(current_section)
        
        return sections


def parse_document(file_path: Union[str, Path], 
                  parser_type: str = 'auto',
                  google_service=None) -> ParsedDocument:
    """Convenience function to parse any supported document type.
    
    Args:
        file_path: Path to document or Google Doc ID
        parser_type: 'auto', 'pdf', 'markdown', 'docx', or 'gdoc'
        google_service: Authenticated Google Docs service (for gdoc)
        
    Returns:
        ParsedDocument with extracted content
    """
    # Initialize parsers
    parsers = {
        'pdf': PDFParser(),
        'markdown': MarkdownParser(),
        'docx': DocxParser(),
        'gdoc': GoogleDocsParser(google_service)
    }
    
    if parser_type == 'auto':
        # Try each parser
        for parser in parsers.values():
            if parser.can_parse(file_path):
                return parser.parse(file_path)
        
        # No parser found
        return ParsedDocument(
            content='',
            error=f"No parser found for file: {file_path}",
            source_path=str(file_path)
        )
    else:
        # Use specific parser
        if parser_type in parsers:
            return parsers[parser_type].parse(file_path)
        else:
            return ParsedDocument(
                content='',
                error=f"Unknown parser type: {parser_type}",
                source_path=str(file_path)
            )