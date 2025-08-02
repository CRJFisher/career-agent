"""Tests for document parsing utilities."""

import pytest
from datetime import datetime
from pathlib import Path
import tempfile
import os
from unittest.mock import Mock, patch, MagicMock
import PyPDF2
from docx import Document
import frontmatter

from utils.document_parser import (
    DocumentSection, ParsedDocument, DocumentParser,
    PDFParser, MarkdownParser, DocxParser, GoogleDocsParser,
    parse_document
)


class TestDocumentSection:
    """Tests for DocumentSection class."""
    
    def test_document_section_creation(self):
        """Test creating DocumentSection instance."""
        section = DocumentSection(
            title="Introduction",
            content="This is the introduction.",
            level=1,
            start_page=1,
            end_page=2
        )
        
        assert section.title == "Introduction"
        assert section.content == "This is the introduction."
        assert section.level == 1
        assert section.start_page == 1
        assert section.end_page == 2
    
    def test_document_section_to_dict(self):
        """Test converting DocumentSection to dictionary."""
        section = DocumentSection(
            title="Methods",
            content="Research methods...",
            level=2
        )
        
        result = section.to_dict()
        
        assert result["title"] == "Methods"
        assert result["content"] == "Research methods..."
        assert result["level"] == 2
        assert result["start_page"] is None
        assert result["end_page"] is None


class TestParsedDocument:
    """Tests for ParsedDocument class."""
    
    def test_parsed_document_creation(self):
        """Test creating ParsedDocument instance."""
        doc = ParsedDocument(
            content="Document content",
            metadata={"author": "John Doe"},
            source_type="pdf",
            source_path="/path/to/doc.pdf"
        )
        
        assert doc.content == "Document content"
        assert doc.metadata == {"author": "John Doe"}
        assert doc.source_type == "pdf"
        assert doc.source_path == "/path/to/doc.pdf"
        assert doc.sections == []
        assert doc.error is None
    
    def test_parsed_document_to_dict(self):
        """Test converting ParsedDocument to dictionary."""
        section = DocumentSection("Title", "Content", 1)
        doc = ParsedDocument(
            content="Full content",
            metadata={"pages": 10},
            sections=[section],
            source_type="docx"
        )
        
        result = doc.to_dict()
        
        assert result["content"] == "Full content"
        assert result["metadata"] == {"pages": 10}
        assert len(result["sections"]) == 1
        assert result["sections"][0]["title"] == "Title"
        assert result["source_type"] == "docx"
    
    def test_get_sections_by_level(self):
        """Test getting sections by heading level."""
        sections = [
            DocumentSection("Chapter 1", "Content 1", 1),
            DocumentSection("Section 1.1", "Content 1.1", 2),
            DocumentSection("Section 1.2", "Content 1.2", 2),
            DocumentSection("Chapter 2", "Content 2", 1),
        ]
        
        doc = ParsedDocument(content="", sections=sections)
        
        level_1 = doc.get_sections_by_level(1)
        assert len(level_1) == 2
        assert level_1[0].title == "Chapter 1"
        assert level_1[1].title == "Chapter 2"
        
        level_2 = doc.get_sections_by_level(2)
        assert len(level_2) == 2
        assert level_2[0].title == "Section 1.1"
    
    def test_search_content(self):
        """Test searching content."""
        doc = ParsedDocument(
            content="This is a test document. Testing search functionality."
        )
        
        # Case insensitive search
        matches = doc.search_content("test")
        assert len(matches) == 2
        assert matches[0]["text"] == "test"
        assert matches[1]["text"] == "Test"
        
        # Case sensitive search
        matches = doc.search_content("Test", case_sensitive=True)
        assert len(matches) == 1
        assert matches[0]["text"] == "Test"


class TestPDFParser:
    """Tests for PDFParser class."""
    
    @pytest.fixture
    def pdf_parser(self):
        """Create PDFParser instance."""
        return PDFParser()
    
    @pytest.fixture
    def mock_pdf_path(self, tmp_path):
        """Create a mock PDF file path."""
        pdf_path = tmp_path / "test.pdf"
        pdf_path.write_text("")  # Create empty file
        return pdf_path
    
    def test_can_parse_pdf(self, pdf_parser, mock_pdf_path):
        """Test PDF file detection."""
        assert pdf_parser.can_parse(mock_pdf_path)
        assert pdf_parser.can_parse(str(mock_pdf_path))
        
        # Non-PDF file
        txt_path = mock_pdf_path.with_suffix('.txt')
        txt_path.write_text("")
        assert not pdf_parser.can_parse(txt_path)
        
        # Non-existent file
        assert not pdf_parser.can_parse("nonexistent.pdf")
    
    @patch('PyPDF2.PdfReader')
    def test_parse_pdf_success(self, mock_reader_class, pdf_parser, mock_pdf_path):
        """Test successful PDF parsing."""
        # Mock PDF reader
        mock_reader = Mock()
        mock_reader_class.return_value = mock_reader
        
        # Mock pages
        mock_page1 = Mock()
        mock_page1.extract_text.return_value = "Page 1 content\nHEADING ONE\nSome text"
        mock_page2 = Mock()
        mock_page2.extract_text.return_value = "Page 2 content"
        
        mock_reader.pages = [mock_page1, mock_page2]
        mock_reader.metadata = {
            '/Title': 'Test Document',
            '/Author': 'Test Author',
            '/CreationDate': Mock(strftime=lambda fmt: '2024-01-01 00:00:00')
        }
        
        # Parse
        result = pdf_parser.parse(mock_pdf_path)
        
        # Check results
        assert result.source_type == 'pdf'
        assert result.source_path == str(mock_pdf_path.absolute())
        assert "Page 1 content" in result.content
        assert "Page 2 content" in result.content
        assert result.metadata['title'] == 'Test Document'
        assert result.metadata['author'] == 'Test Author'
        assert result.metadata['page_count'] == 2
        assert result.error is None
        
        # Check sections
        assert len(result.sections) > 0
        heading_section = next((s for s in result.sections if s.title == "HEADING ONE"), None)
        assert heading_section is not None
        assert heading_section.level == 1
    
    def test_parse_pdf_error(self, pdf_parser, mock_pdf_path):
        """Test PDF parsing with error."""
        # Make file unreadable
        mock_pdf_path.chmod(0o000)
        
        result = pdf_parser.parse(mock_pdf_path)
        
        assert result.error is not None
        assert result.content == ''
        
        # Restore permissions
        mock_pdf_path.chmod(0o644)


class TestMarkdownParser:
    """Tests for MarkdownParser class."""
    
    @pytest.fixture
    def md_parser(self):
        """Create MarkdownParser instance."""
        return MarkdownParser()
    
    @pytest.fixture
    def mock_md_path(self, tmp_path):
        """Create a mock Markdown file."""
        md_path = tmp_path / "test.md"
        content = """---
title: Test Document
author: John Doe
date: 2024-01-01
---

# Introduction

This is the introduction section.

## Background

Some background information.

### Details

More detailed information here.

# Conclusion

Final thoughts.
"""
        md_path.write_text(content)
        return md_path
    
    def test_can_parse_markdown(self, md_parser, mock_md_path):
        """Test Markdown file detection."""
        assert md_parser.can_parse(mock_md_path)
        
        # Also test .markdown extension
        markdown_path = mock_md_path.with_suffix('.markdown')
        markdown_path.write_text("")
        assert md_parser.can_parse(markdown_path)
        
        # Non-markdown file
        assert not md_parser.can_parse("test.txt")
    
    def test_parse_markdown_success(self, md_parser, mock_md_path):
        """Test successful Markdown parsing."""
        result = md_parser.parse(mock_md_path)
        
        # Check basic properties
        assert result.source_type == 'markdown'
        assert result.source_path == str(mock_md_path.absolute())
        assert "This is the introduction section." in result.content
        assert result.error is None
        
        # Check metadata from frontmatter
        assert result.metadata['title'] == 'Test Document'
        assert result.metadata['author'] == 'John Doe'
        assert result.metadata['date'] == '2024-01-01'
        assert 'file_modified' in result.metadata
        assert 'file_size' in result.metadata
        
        # Check sections
        assert len(result.sections) == 4
        assert result.sections[0].title == "Introduction"
        assert result.sections[0].level == 1
        assert result.sections[1].title == "Background"
        assert result.sections[1].level == 2
        assert result.sections[2].title == "Details"
        assert result.sections[2].level == 3
        assert result.sections[3].title == "Conclusion"
        assert result.sections[3].level == 1
        
        # Check raw data
        assert 'markdown' in result.raw_data
        assert 'html' in result.raw_data
    
    def test_parse_markdown_no_frontmatter(self, md_parser, tmp_path):
        """Test parsing Markdown without frontmatter."""
        md_path = tmp_path / "simple.md"
        md_path.write_text("# Title\n\nJust content.")
        
        result = md_parser.parse(md_path)
        
        assert result.content == "# Title\n\nJust content."
        assert len(result.sections) == 1
        assert result.sections[0].title == "Title"


class TestDocxParser:
    """Tests for DocxParser class."""
    
    @pytest.fixture
    def docx_parser(self):
        """Create DocxParser instance."""
        return DocxParser()
    
    def test_can_parse_docx(self, docx_parser, tmp_path):
        """Test Word document detection."""
        docx_path = tmp_path / "test.docx"
        docx_path.write_text("")
        assert docx_parser.can_parse(docx_path)
        
        doc_path = tmp_path / "test.doc"
        doc_path.write_text("")
        assert docx_parser.can_parse(doc_path)
        
        assert not docx_parser.can_parse("test.txt")
    
    @patch('docx.Document')
    def test_parse_docx_success(self, mock_document_class, docx_parser, tmp_path):
        """Test successful DOCX parsing."""
        docx_path = tmp_path / "test.docx"
        docx_path.write_text("")
        
        # Mock document
        mock_doc = Mock()
        mock_document_class.return_value = mock_doc
        
        # Mock paragraphs
        mock_para1 = Mock()
        mock_para1.text = "Document Title"
        mock_para1.style.name = "Heading 1"
        
        mock_para2 = Mock()
        mock_para2.text = "This is paragraph text."
        mock_para2.style.name = "Normal"
        
        mock_para3 = Mock()
        mock_para3.text = "Section Title"
        mock_para3.style.name = "Heading 2"
        
        mock_doc.paragraphs = [mock_para1, mock_para2, mock_para3]
        mock_doc.tables = []
        
        # Mock properties
        mock_props = Mock()
        mock_props.author = "Test Author"
        mock_props.title = "Test Title"
        mock_props.subject = None
        mock_props.created = datetime(2024, 1, 1)
        mock_props.modified = datetime(2024, 1, 2)
        mock_doc.core_properties = mock_props
        
        # Parse
        result = docx_parser.parse(docx_path)
        
        # Check results
        assert result.source_type == 'docx'
        assert "Document Title" in result.content
        assert "This is paragraph text." in result.content
        assert result.metadata['author'] == "Test Author"
        assert result.metadata['title'] == "Test Title"
        assert len(result.sections) == 2
        assert result.sections[0].title == "Document Title"
        assert result.sections[0].level == 1
        assert result.error is None
    
    def test_parse_doc_format_error(self, docx_parser, tmp_path):
        """Test parsing legacy .doc format."""
        doc_path = tmp_path / "old.doc"
        doc_path.write_text("")
        
        result = docx_parser.parse(doc_path)
        
        assert result.error == "Legacy .doc format not supported. Please convert to .docx"
        assert result.content == ''


class TestGoogleDocsParser:
    """Tests for GoogleDocsParser class."""
    
    @pytest.fixture
    def mock_service(self):
        """Create mock Google Docs service."""
        return Mock()
    
    @pytest.fixture
    def gdoc_parser(self, mock_service):
        """Create GoogleDocsParser instance."""
        return GoogleDocsParser(mock_service)
    
    def test_can_parse_gdoc(self, gdoc_parser):
        """Test Google Doc ID detection."""
        # Valid ID format
        assert gdoc_parser.can_parse("1234567890abcdefghijklmnopqrstuvwxyz")
        
        # Too short
        assert not gdoc_parser.can_parse("short123")
        
        # Invalid characters
        assert not gdoc_parser.can_parse("123@456#789")
        
        # Path-like string
        assert not gdoc_parser.can_parse("/path/to/file")
    
    def test_parse_gdoc_success(self, gdoc_parser, mock_service):
        """Test successful Google Doc parsing."""
        doc_id = "1234567890abcdefghijklmnopqrstuvwxyz"
        
        # Mock document structure
        mock_gdoc = {
            'title': 'Test Google Doc',
            'body': {
                'content': [
                    {
                        'paragraph': {
                            'elements': [
                                {'textRun': {'content': 'Main Heading\n'}},
                            ],
                            'paragraphStyle': {'namedStyleType': 'HEADING_1'}
                        }
                    },
                    {
                        'paragraph': {
                            'elements': [
                                {'textRun': {'content': 'This is paragraph text.\n'}},
                            ],
                            'paragraphStyle': {'namedStyleType': 'NORMAL_TEXT'}
                        }
                    }
                ]
            }
        }
        
        mock_service.documents().get().execute.return_value = mock_gdoc
        
        # Parse
        result = gdoc_parser.parse(doc_id)
        
        # Check results
        assert result.source_type == 'gdoc'
        assert result.source_path == doc_id
        assert "Main Heading" in result.content
        assert "This is paragraph text." in result.content
        assert result.metadata['title'] == 'Test Google Doc'
        assert len(result.sections) == 1
        assert result.sections[0].title == "Main Heading"
        assert result.sections[0].level == 1
        assert result.error is None
    
    def test_parse_gdoc_no_service(self):
        """Test parsing without service."""
        parser = GoogleDocsParser()  # No service
        result = parser.parse("123456")
        
        assert result.error == "Google Docs service not initialized"
        assert result.content == ''


class TestParseDocument:
    """Tests for parse_document convenience function."""
    
    @patch('utils.document_parser.PDFParser.parse')
    @patch('utils.document_parser.PDFParser.can_parse')
    def test_parse_document_auto_pdf(self, mock_can_parse, mock_parse):
        """Test auto-detection of PDF."""
        mock_can_parse.return_value = True
        mock_parse.return_value = ParsedDocument(content="PDF content", source_type="pdf")
        
        result = parse_document("test.pdf")
        
        assert result.content == "PDF content"
        assert result.source_type == "pdf"
        mock_can_parse.assert_called_once()
        mock_parse.assert_called_once()
    
    @patch('utils.document_parser.MarkdownParser.parse')
    def test_parse_document_specific_parser(self, mock_parse):
        """Test using specific parser."""
        mock_parse.return_value = ParsedDocument(content="MD content", source_type="markdown")
        
        result = parse_document("test.md", parser_type="markdown")
        
        assert result.content == "MD content"
        assert result.source_type == "markdown"
        mock_parse.assert_called_once()
    
    def test_parse_document_no_parser(self):
        """Test when no parser is found."""
        result = parse_document("unknown.xyz")
        
        assert result.error == "No parser found for file: unknown.xyz"
        assert result.content == ''
    
    def test_parse_document_unknown_parser_type(self):
        """Test with unknown parser type."""
        result = parse_document("test.pdf", parser_type="unknown")
        
        assert result.error == "Unknown parser type: unknown"
        assert result.content == ''